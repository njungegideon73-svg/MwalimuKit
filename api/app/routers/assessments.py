"""Assessment generation + CRUD + export."""
import io
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import StreamingResponse
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.ai.factory import get_provider
from app.core.db import get_db
from app.core.deps import CurrentUser
from app.core.rate_limit import rate_limit_generate
from app.models.assessment import Assessment
from app.models.curriculum import LearningArea
from app.models.prompt_history import PromptHistory
from app.schemas.assessment import (
    AssessmentIn,
    AssessmentOut,
    AssessmentUpdate,
    GenerateAssessmentRequest,
    GenerateAssessmentResponse,
)
from app.services.assessments import (
    create_assessment as svc_create,
)
from app.services.assessments import (
    duplicate_assessment as svc_duplicate,
)
from app.services.assessments import (
    get_assessment as svc_get,
)
from app.services.assessments import (
    get_assessment_or_404,
)
from app.services.assessments import (
    list_assessments as svc_list,
)
from app.services.assessments import (
    soft_delete as svc_soft_delete,
)
from app.services.assessments import (
    toggle_favourite as svc_toggle_favourite,
)
from app.services.assessments import (
    update_assessment as svc_update,
)
from app.utils.activity_logger import log_activity

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


router = APIRouter()


@router.post("/generate", response_model=GenerateAssessmentResponse)
async def generate(
    request: Request,
    req: GenerateAssessmentRequest,
    user: CurrentUser,
    db: AsyncSession = Depends(get_db),
    _rate: None = Depends(rate_limit_generate),
) -> GenerateAssessmentResponse:
    provider = get_provider()
    result = await provider.generate_assessment(
        learning_area=req.learning_area_code,
        strand=req.strand_code,
        sub_strand=", ".join(req.sub_strand_codes),
        grade_level=req.grade_level,
        teacher_prompt=req.teacher_prompt,
        item_count=req.item_count,
        include_diagrams=req.include_diagrams,
    )

    history = PromptHistory(
        user_id=user.id,
        school_id=user.school_id,
        learning_area_code=req.learning_area_code,
        strand_code=req.strand_code,
        sub_strand_codes=req.sub_strand_codes,
        grade_level=req.grade_level,
        teacher_prompt=req.teacher_prompt,
        item_count=req.item_count,
        response_rubric=result.rubric.model_dump() if hasattr(result.rubric, "model_dump") else result.rubric,
        response_items=[i.model_dump() for i in result.items] if hasattr(result.items[0], "model_dump") else result.items,
        provider=result.provider,
        model=result.model,
    )
    db.add(history)
    await db.commit()

    return GenerateAssessmentResponse(
        rubric=result.rubric,
        items=result.items,
        provider=result.provider,
        model=result.model,
    )


@router.get("", response_model=list[AssessmentOut])
async def list_assessments(user: CurrentUser, db: AsyncSession = Depends(get_db)) -> list[AssessmentOut]:
    return await svc_list(db, user.school_id)


@router.post("", response_model=AssessmentOut)
async def create(
    payload: AssessmentIn, user: CurrentUser, db: AsyncSession = Depends(get_db)
) -> AssessmentOut:
    result = await svc_create(db, payload, user.school_id, user.id)
    await log_activity(
        db,
        user_id=user.id,
        school_id=user.school_id,
        action="assessment.created",
        details={"name": result.name, "learning_area_code": payload.learning_area_code},
    )
    return result


@router.get("/{assessment_id}", response_model=AssessmentOut)
async def get(
    assessment_id: UUID, user: CurrentUser, db: AsyncSession = Depends(get_db)
) -> AssessmentOut:
    return await svc_get(db, assessment_id, user.school_id)


@router.post("/{assessment_id}/duplicate", response_model=AssessmentOut)
async def duplicate(
    assessment_id: UUID, user: CurrentUser, db: AsyncSession = Depends(get_db)
) -> AssessmentOut:
    result = await svc_duplicate(db, assessment_id, user.school_id, user.id)
    await log_activity(
        db,
        user_id=user.id,
        school_id=user.school_id,
        action="assessment.duplicated",
        details={"name": result.name},
    )
    return result


@router.patch("/{assessment_id}", response_model=AssessmentOut)
async def update(
    assessment_id: UUID,
    payload: AssessmentUpdate,
    user: CurrentUser,
    db: AsyncSession = Depends(get_db),
) -> AssessmentOut:
    result = await svc_update(db, assessment_id, user.school_id, payload)
    await log_activity(
        db,
        user_id=user.id,
        school_id=user.school_id,
        action="assessment.updated",
        details={"name": result.name},
    )
    return result


@router.post("/{assessment_id}/favourite", response_model=AssessmentOut)
async def toggle_favourite(
    assessment_id: UUID,
    user: CurrentUser,
    db: AsyncSession = Depends(get_db),
) -> AssessmentOut:
    result = await svc_toggle_favourite(db, assessment_id, user.school_id)
    await log_activity(
        db,
        user_id=user.id,
        school_id=user.school_id,
        action="assessment.favourited",
        details={"name": result.name, "is_favourite": result.is_favourite},
    )
    return result


@router.delete("/{assessment_id}")
async def soft_delete(
    assessment_id: UUID, user: CurrentUser, db: AsyncSession = Depends(get_db)
) -> dict:
    a = await get_assessment_or_404(db, assessment_id, user.school_id)
    await log_activity(
        db,
        user_id=user.id,
        school_id=user.school_id,
        action="assessment.deleted",
        details={"name": a.name},
    )
    return await svc_soft_delete(db, assessment_id, user.school_id)


async def _load_for_export(db: AsyncSession, assessment_id: UUID, school_id: UUID) -> tuple[Assessment, str]:
    a = (
        await db.execute(
            select(Assessment).where(
                Assessment.id == assessment_id,
                Assessment.school_id == school_id,
                Assessment.deleted_at.is_(None),
            )
        )
    ).scalar_one_or_none()
    if a is None:
        raise HTTPException(status_code=404, detail="Assessment not found")

    la = (
        await db.execute(select(LearningArea).where(LearningArea.id == a.learning_area_id))
    ).scalar_one_or_none()
    la_name = la.name if la else "Unknown"
    return a, la_name


@router.get("/{assessment_id}/export/pdf")
async def export_assessment_pdf(
    assessment_id: UUID,
    user: CurrentUser,
    db: AsyncSession = Depends(get_db),
    mode: str = Query(default="questions"),
):
    if mode not in ("questions", "answer-key"):
        raise HTTPException(status_code=400, detail="Invalid export mode. Use 'questions' or 'answer-key'.")

    a, la_name = await _load_for_export(db, assessment_id, user.school_id)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    elements: list = []

    title_style = ParagraphStyle("Title2", parent=styles["Title"], fontSize=18, spaceAfter=6)
    subtitle_style = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=11, spaceAfter=4)
    body_style = styles["Normal"]
    small_style = ParagraphStyle("Small", parent=body_style, fontSize=9, spaceAfter=2)

    elements.append(Paragraph(f"Assessment: {a.name}", title_style))
    elements.append(Paragraph(f"Learning Area: {la_name}", subtitle_style))
    elements.append(Paragraph(f"Strand: {a.strand_code or '-'}", subtitle_style))
    elements.append(Paragraph(f"Source: {a.source.value if hasattr(a.source, 'value') else a.source}", subtitle_style))
    elements.append(Spacer(1, 0.5 * cm))

    total_marks = 0
    for idx, item in enumerate(a.items or [], start=1):
        item_max = item.get("max_level", 4) if isinstance(item.get("max_level", 4), (int, float)) else 4
        total_marks += item_max

        elements.append(Paragraph(f"<b>Question {idx}</b>", body_style))
        elements.append(Paragraph(item.get("stem", ""), body_style))
        if item.get("diagram_description"):
            elements.append(Paragraph(
                f"<b>Diagram / Visual:</b> {item['diagram_description']}",
                ParagraphStyle("Diagram", parent=body_style, fontSize=10, textColor=colors.HexColor("#555555"), spaceAfter=4)
            ))
        if mode == "questions":
            elements.append(Paragraph(f"<b>Answer guide:</b> {item.get('answer_guide', '') or 'N/A'}", small_style))
        elements.append(Paragraph(f"<b>Max marks:</b> {item_max}", small_style))
        elements.append(Spacer(1, 0.3 * cm))

    elements.append(Spacer(1, 0.3 * cm))
    elements.append(Paragraph(f"<b>Total Marks: {total_marks}</b>", ParagraphStyle("Total", parent=body_style, fontSize=12, spaceAfter=6)))

    if mode == "answer-key":
        elements.append(Paragraph("<b>Answer Key</b>", ParagraphStyle("KeyTitle", parent=body_style, fontSize=14, spaceAfter=6)))
        for idx, item in enumerate(a.items or [], start=1):
            elements.append(Paragraph(
                f"{idx}. <b>{item.get('answer_guide', '') or 'N/A'}</b>  (max {item.get('max_level', 4)} marks, criterion: {item.get('criterion', 'N/A')})",
                small_style
            ))

    if a.rubric:
        elements.append(Spacer(1, 0.3 * cm))
        elements.append(Paragraph("<b>Rubric</b>", ParagraphStyle("RubricTitle", parent=body_style, fontSize=14, spaceAfter=6)))
        for level in a.rubric.get("levels", []):
            elements.append(Paragraph(
                f"Level {level.get('level')}: <b>{level.get('label', '')}</b> — {level.get('descriptor', '')}",
                small_style
            ))
        if a.rubric.get("criteria"):
            elements.append(Spacer(1, 0.2 * cm))
            for crit in a.rubric.get("criteria", []):
                elements.append(Paragraph(
                    f"Criterion: {crit.get('label', '')} ({crit.get('id', '')})",
                    small_style
                ))

    doc.build(elements)
    buffer.seek(0)
    mode_suffix = "-answer-key" if mode == "answer-key" else ""
    filename = f"assessment-{a.name.replace(' ', '-').lower()}{mode_suffix}.pdf"
    return StreamingResponse(buffer, media_type="application/pdf", headers={
        "Content-Disposition": f"attachment; filename=\"{filename}\""
    })


@router.get("/{assessment_id}/export/docx")
async def export_assessment_docx(
    assessment_id: UUID,
    user: CurrentUser,
    db: AsyncSession = Depends(get_db),
    mode: str = Query(default="questions"),
):
    if mode not in ("questions", "answer-key"):
        raise HTTPException(status_code=400, detail="Invalid export mode. Use 'questions' or 'answer-key'.")
    if not HAS_DOCX:
        raise HTTPException(status_code=501, detail="DOCX export is not available on this server")

    a, la_name = await _load_for_export(db, assessment_id, user.school_id)

    document = Document()
    document.add_heading(f"Assessment: {a.name}", level=1)
    document.add_paragraph(f"Learning Area: {la_name}")
    document.add_paragraph(f"Strand: {a.strand_code or '-'}")
    document.add_paragraph(f"Source: {a.source.value if hasattr(a.source, 'value') else a.source}")

    total_marks = 0
    for idx, item in enumerate(a.items or [], start=1):
        item_max = item.get("max_level", 4)
        total_marks += item_max
        p = document.add_paragraph()
        p.add_run(f"Question {idx}: ").bold = True
        p.add_run(item.get("stem", ""))
        p.add_run(f"  [Max marks: {item_max}]")
        if item.get("diagram_description"):
            d = document.add_paragraph()
            d.add_run("Diagram / Visual: ").bold = True
            d.add_run(item["diagram_description"])
        if mode == "questions":
            document.add_paragraph(f"Answer guide: {item.get('answer_guide', '') or 'N/A'}")

    document.add_paragraph(f"Total Marks: {total_marks}")

    if mode == "answer-key":
        document.add_heading("Answer Key", level=2)
        for idx, item in enumerate(a.items or [], start=1):
            document.add_paragraph(
                f"{idx}. {item.get('answer_guide', '') or 'N/A'}  (max {item.get('max_level', 4)} marks, criterion: {item.get('criterion', 'N/A')})"
            )

    if a.rubric:
        document.add_heading("Rubric", level=2)
        for level in a.rubric.get("levels", []):
            p = document.add_paragraph(style='List Bullet')
            p.add_run(f"Level {level.get('level')}: {level.get('label', '')} — {level.get('descriptor', '')}")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    mode_suffix = "-answer-key" if mode == "answer-key" else ""
    filename = f"assessment-{a.name.replace(' ', '-').lower()}{mode_suffix}.docx"
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
        "Content-Disposition": f"attachment; filename=\"{filename}\""
    })
