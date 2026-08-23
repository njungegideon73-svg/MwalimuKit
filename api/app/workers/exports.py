"""Background workers for export jobs."""
from __future__ import annotations

import asyncio
import io
from uuid import UUID

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import select
from sqlalchemy.ext.asyncio import async_sessionmaker

from app.core.db import engine
from app.core.jobs import job_expires_at
from app.models.job import Job, JobStatus
import dramatiq


async def _get_job(session, job_id: UUID) -> Job | None:
    return await session.get(Job, job_id)


async def _update_job(session, job: Job, **kwargs) -> None:
    for k, v in kwargs.items():
        setattr(job, k, v)
    job.updated_at = job_expires_at()
    session.add(job)
    await session.commit()


@dramatiq.actor
def process_export_job(job_id: str) -> None:
    asyncio.run(_process_export_job_async(UUID(job_id)))


async def _process_export_job_async(job_id: UUID) -> None:

    async_session = async_sessionmaker(engine, expire_on_commit=False)

    async with async_session() as session:
        job = await _get_job(session, job_id)
        if job is None:
            return

        job.status = JobStatus.processing
        await session.commit()

        try:
            if job.type == "assessment_pdf":
                data, result = await _build_assessment_pdf(session, job.payload)
            elif job.type == "assessment_docx":
                data, result = await _build_assessment_docx(session, job.payload)
            elif job.type == "report_card_pdf":
                data, result = await _build_report_card_pdf(session, job.payload)
            elif job.type == "sba_report_card_pdf":
                data, result = await _build_sba_report_card_pdf(session, job.payload)
            elif job.type == "class_summary_csv":
                data, result = await _build_class_summary_csv(session, job.payload)
            elif job.type == "term_exam_class_csv":
                data, result = await _build_term_exam_class_csv(session, job.payload)
            else:
                raise ValueError(f"Unknown job type: {job.type}")

            job.status = JobStatus.completed
            job.result = result
            job.file_data = data
            job.expires_at = job_expires_at()
            await session.commit()

        except Exception as exc:
            job.status = JobStatus.failed
            job.error = str(exc)
            await session.commit()
            raise


async def _get_school_name(session, school_id: UUID) -> str:
    from app.models.school import School
    school = await session.get(School, school_id)
    return school.name if school else "School"


async def _get_learning_area_name(session, learning_area_id: UUID) -> str:
    from app.models.curriculum import LearningArea
    la = await session.get(LearningArea, learning_area_id)
    return la.name if la else "Unknown"


def _make_result(data: bytes, filename: str, content_type: str) -> dict:
    return {
        "filename": filename,
        "size_bytes": len(data),
        "content_type": content_type,
    }


async def _build_assessment_pdf(session, payload: dict) -> dict:
    from app.models.assessment import Assessment

    assessment_id = UUID(payload["assessment_id"])
    school_id = UUID(payload["school_id"])
    mode = payload.get("mode", "questions")

    assessment = (
        await session.execute(
            select(Assessment).where(
                Assessment.id == assessment_id,
                Assessment.school_id == school_id,
                Assessment.deleted_at.is_(None),
            )
        )
    ).scalar_one_or_none()
    if assessment is None:
        raise ValueError("Assessment not found")

    la_name = await _get_learning_area_name(session, assessment.learning_area_id)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    elements = []

    title_style = ParagraphStyle("Title2", parent=styles["Title"], fontSize=18, spaceAfter=6)
    subtitle_style = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=11, spaceAfter=4)
    body_style = styles["Normal"]
    small_style = ParagraphStyle("Small", parent=body_style, fontSize=9, spaceAfter=2)

    elements.append(Paragraph(f"Assessment: {assessment.name}", title_style))
    elements.append(Paragraph(f"Learning Area: {la_name}", subtitle_style))
    elements.append(Paragraph(f"Strand: {assessment.strand_code or '-'}", subtitle_style))
    elements.append(Paragraph(f"Source: {assessment.source.value if hasattr(assessment.source, 'value') else assessment.source}", subtitle_style))
    elements.append(Spacer(1, 0.5 * cm))

    total_marks = 0
    for idx, item in enumerate(assessment.items or [], start=1):
        item_max = item.get("max_level", 4) if isinstance(item.get("max_level", 4), (int, float)) else 4
        total_marks += item_max
        elements.append(Paragraph(f"<b>Question {idx}</b>", body_style))
        elements.append(Paragraph(item.get("stem", ""), body_style))
        if item.get("diagram_description"):
            elements.append(Paragraph(
                f"<b>Diagram / Visual:</b> {item['diagram_description']}",
                ParagraphStyle("Diagram", parent=body_style, fontSize=10, textColor=colors.HexColor("#555555"), spaceAfter=4)
            ))
        if mode == "questions":
            elements.append(Paragraph(f"<b>Answer guide:</b> {item.get('answer_guide', '') or 'N/A'}", small_style))
        elements.append(Paragraph(f"<b>Max marks:</b> {item_max}", small_style))
        elements.append(Spacer(1, 0.3 * cm))

    elements.append(Spacer(1, 0.3 * cm))
    elements.append(Paragraph(f"<b>Total Marks: {total_marks}</b>", ParagraphStyle("Total", parent=body_style, fontSize=12, spaceAfter=6)))

    if mode == "answer-key":
        elements.append(Paragraph("<b>Answer Key</b>", ParagraphStyle("KeyTitle", parent=body_style, fontSize=14, spaceAfter=6)))
        for idx, item in enumerate(assessment.items or [], start=1):
            elements.append(Paragraph(
                f"{idx}. <b>{item.get('answer_guide', '') or 'N/A'}</b>  (max {item.get('max_level', 4)} marks, criterion: {item.get('criterion', 'N/A')})",
                small_style
            ))

    if assessment.rubric:
        elements.append(Spacer(1, 0.3 * cm))
        elements.append(Paragraph("<b>Rubric</b>", ParagraphStyle("RubricTitle", parent=body_style, fontSize=14, spaceAfter=6)))
        for level in assessment.rubric.get("levels", []):
            elements.append(Paragraph(
                f"Level {level.get('level')}: <b>{level.get('label', '')}</b> — {level.get('descriptor', '')}",
                small_style
            ))
        if assessment.rubric.get("criteria"):
            elements.append(Spacer(1, 0.2 * cm))
            for crit in assessment.rubric.get("criteria", []):
                elements.append(Paragraph(
                    f"Criterion: {crit.get('label', '')} ({crit.get('id', '')})",
                    small_style
                ))

    doc.build(elements)
    buffer.seek(0)
    data = buffer.read()
    mode_suffix = "-answer-key" if mode == "answer-key" else ""
    filename = f"assessment-{assessment.name.replace(' ', '-').lower()}{mode_suffix}.pdf"
    result = _make_result(data, filename, "application/pdf")
    return data, result


async def _build_assessment_docx(session, payload: dict) -> dict:
    from app.models.assessment import Assessment

    assessment_id = UUID(payload["assessment_id"])
    school_id = UUID(payload["school_id"])
    mode = payload.get("mode", "questions")

    assessment = (
        await session.execute(
            select(Assessment).where(
                Assessment.id == assessment_id,
                Assessment.school_id == school_id,
                Assessment.deleted_at.is_(None),
            )
        )
    ).scalar_one_or_none()
    if assessment is None:
        raise ValueError("Assessment not found")

    document = Document()
    document.add_heading(f"Assessment: {assessment.name}", level=1)
    document.add_paragraph(f"Learning Area: {await _get_learning_area_name(session, assessment.learning_area_id)}")
    document.add_paragraph(f"Strand: {assessment.strand_code or '-'}")
    document.add_paragraph(f"Source: {assessment.source.value if hasattr(assessment.source, 'value') else assessment.source}")

    total_marks = 0
    for idx, item in enumerate(assessment.items or [], start=1):
        item_max = item.get("max_level", 4)
        total_marks += item_max
        p = document.add_paragraph()
        p.add_run(f"Question {idx}: ").bold = True
        p.add_run(item.get("stem", ""))
        p.add_run(f"  [Max marks: {item_max}]")
        if item.get("diagram_description"):
            d = document.add_paragraph()
            d.add_run("Diagram / Visual: ").bold = True
            d.add_run(item["diagram_description"])
        if mode == "questions":
            document.add_paragraph(f"Answer guide: {item.get('answer_guide', '') or 'N/A'}")

    document.add_paragraph(f"Total Marks: {total_marks}")

    if mode == "answer-key":
        document.add_heading("Answer Key", level=2)
        for idx, item in enumerate(assessment.items or [], start=1):
            document.add_paragraph(
                f"{idx}. {item.get('answer_guide', '') or 'N/A'}  (max {item.get('max_level', 4)} marks, criterion: {item.get('criterion', 'N/A')})"
            )

    if assessment.rubric:
        document.add_heading("Rubric", level=2)
        for level in assessment.rubric.get("levels", []):
            p = document.add_paragraph(style='List Bullet')
            p.add_run(f"Level {level.get('level')}: {level.get('label', '')} — {level.get('descriptor', '')}")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    data = buffer.read()
    mode_suffix = "-answer-key" if mode == "answer-key" else ""
    filename = f"assessment-{assessment.name.replace(' ', '-').lower()}{mode_suffix}.docx"
    result = _make_result(data, filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    return data, result


async def _build_report_card_pdf(session, payload: dict) -> dict:
    from app.models.assessment import Assessment
    from app.models.learner import Learner
    from app.models.run import AssessmentRun
    from app.models.school import School
    from app.models.school_class import SchoolClass
    from app.models.score import Score

    learner_id = UUID(payload["learner_id"])
    run_id = UUID(payload["run_id"])
    school_id = UUID(payload["school_id"])

    learner = (
        await session.execute(
            select(Learner).where(
                Learner.id == learner_id,
                Learner.school_id == school_id,
                Learner.deleted_at.is_(None),
            )
        )
    ).scalar_one_or_none()
    if learner is None:
        raise ValueError("Learner not found")

    run = (
        await session.execute(
            select(AssessmentRun).where(AssessmentRun.id == run_id, AssessmentRun.school_id == school_id)
        )
    ).scalar_one_or_none()
    if run is None:
        raise ValueError("Run not found")

    assessment = (
        await session.execute(
            select(Assessment).where(
                Assessment.id == run.assessment_id,
                Assessment.school_id == school_id,
                Assessment.deleted_at.is_(None),
            )
        )
    ).scalar_one_or_none()
    if assessment is None:
        raise ValueError("Assessment not found")

    school_class = (
        await session.execute(
            select(SchoolClass).where(SchoolClass.id == run.class_id, SchoolClass.school_id == school_id)
        )
    ).scalar_one_or_none()
    if school_class is None:
        raise ValueError("Class not found")

    if school_class.id != learner.class_id:
        raise ValueError("Learner not in this class")

    school = await session.get(School, school_id)

    scores = (
        await session.execute(
            select(Score).where(Score.run_id == run_id, Score.learner_id == learner_id).order_by(Score.item_id)
        )
    ).scalars().all()
    if not scores:
        raise ValueError("No scores found for this learner in this run")

    def _rubric_label(assessment_obj, level):
        if level is None:
            return "N/A"
        rubric = assessment_obj.rubric or {}
        for rl in rubric.get("levels", []):
            if rl.get("level") == level:
                return rl.get("label", str(level))
        return str(level)

    def _item_title(assessment_obj, item_id):
        for item in assessment_obj.items or []:
            if item.get("id") == item_id:
                return item.get("stem", item_id)
        return item_id

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    elements = []

    title_style = ParagraphStyle("Title2", parent=styles["Title"], fontSize=16, spaceAfter=6)
    subtitle_style = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=11, alignment=1, spaceAfter=4)
    body_style = styles["Normal"]
    small_style = ParagraphStyle("Small", parent=body_style, fontSize=9, spaceAfter=2)

    school_name = school.name if school else "School"
    school_code = school.code if school else ""
    elements.append(Paragraph(f"{school_name}", title_style))
    elements.append(Paragraph(f"School Code: {school_code}", subtitle_style))
    elements.append(Spacer(1, 0.4 * cm))

    gender_display = (learner.gender or "-").upper()
    info_data = [
        [Paragraph("<b>Learner Name:</b>", body_style), Paragraph(learner.full_name, body_style),
         Paragraph("<b>Admission No:</b>", body_style), Paragraph(learner.admission_no or "-", body_style)],
        [Paragraph("<b>Gender:</b>", body_style), Paragraph(gender_display, body_style),
         Paragraph("<b>Class:</b>", body_style), Paragraph(school_class.name, body_style)],
    ]
    info_table = Table(info_data, colWidths=[3.5 * cm, 5.5 * cm, 3.5 * cm, 5.5 * cm])
    info_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 0.3 * cm))

    elements.append(Paragraph(f"<b>Assessment:</b> {assessment.name}", body_style))
    elements.append(Paragraph(f"<b>Term:</b> {run.term or '-'}", body_style))
    elements.append(Spacer(1, 0.5 * cm))

    header = [
        Paragraph("<b>#</b>", body_style),
        Paragraph("<b>Item</b>", body_style),
        Paragraph("<b>Level</b>", body_style),
        Paragraph("<b>Rubric Label</b>", body_style),
    ]
    table_data = [header]

    max_possible = 0
    total_scored = 0
    for idx, score in enumerate(scores, start=1):
        item_stem = _item_title(assessment, score.item_id)
        level = score.level
        label = _rubric_label(assessment, level)
        item_max = 4
        for item in assessment.items or []:
            if item.get("id") == score.item_id:
                item_max = item.get("max_level", 4)
                break
        max_possible += item_max
        total_scored += level if level else 0
        table_data.append([
            Paragraph(str(idx), body_style),
            Paragraph(item_stem[:80], small_style),
            Paragraph(str(level) if level is not None else "-", body_style),
            Paragraph(label, body_style),
        ])

    percentage = round((total_scored / max_possible) * 100, 1) if max_possible > 0 else 0.0
    col_widths = [1.2 * cm, 9.5 * cm, 2 * cm, 5 * cm]
    scores_table = Table(table_data, colWidths=col_widths, repeatRows=1)
    scores_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
    ]))
    elements.append(scores_table)
    elements.append(Spacer(1, 0.5 * cm))

    pct_style = ParagraphStyle("Pct", parent=body_style, fontSize=12, alignment=1, spaceAfter=6)
    elements.append(Paragraph(
        f"<b>Total Score: {total_scored} / {max_possible}  ({percentage}%)</b>",
        pct_style,
    ))
    elements.append(Spacer(1, 1.5 * cm))

    sig_style = ParagraphStyle("Sig", parent=body_style, fontSize=10)
    elements.append(Paragraph("_" * 40 + "          " + "_" * 20, sig_style))
    elements.append(Paragraph("Teacher Signature                                   Date", sig_style))

    doc.build(elements)
    buffer.seek(0)
    data = buffer.read()
    filename = f"report_card_{learner.full_name.replace(' ', '_')}_{run_id}.pdf"
    result = _make_result(data, filename, "application/pdf")
    return data, result


async def _build_sba_report_card_pdf(session, payload: dict) -> dict:
    from app.models.curriculum import LearningArea
    from app.models.learner import Learner
    from app.models.learner_exam_score import LearnerExamScore
    from app.models.school import School
    from app.models.school_class import SchoolClass
    from app.models.term_exam import TermExam

    learner_id = UUID(payload["learner_id"])
    academic_year = payload["academic_year"]
    school_id = UUID(payload["school_id"])

    learner = (
        await session.execute(
            select(Learner).where(
                Learner.id == learner_id,
                Learner.school_id == school_id,
                Learner.deleted_at.is_(None),
            )
        )
    ).scalar_one_or_none()
    if learner is None:
        raise ValueError("Learner not found")

    school_class = (
        await session.execute(
            select(SchoolClass).where(SchoolClass.id == learner.class_id, SchoolClass.school_id == school_id)
        )
    ).scalar_one_or_none()
    if school_class is None:
        raise ValueError("Class not found")

    school = await session.get(School, school_id)

    exams = (
        await session.execute(
            select(TermExam).where(
                TermExam.class_id == learner.class_id,
                TermExam.academic_year == academic_year,
                TermExam.school_id == school_id,
            ).order_by(TermExam.term, TermExam.exam_type)
        )
    ).scalars().all()

    exam_ids = [e.id for e in exams]
    if not exam_ids:
        raise ValueError("No term exams found for this learner in this academic year")

    scores = (
        await session.execute(
            select(LearnerExamScore).where(
                LearnerExamScore.learner_id == learner_id,
                LearnerExamScore.term_exam_id.in_(exam_ids),
            )
        )
    ).scalars().all()
    score_map = {str(s.term_exam_id): s for s in scores}

    subjects = []
    exam_type_scores = {}
    for exam in exams:
        la = await session.get(LearningArea, exam.learning_area_id)
        subject_name = la.name if la else "Unknown"
        score = score_map.get(str(exam.id))
        marks = score.marks if score else None
        pct = round((marks / exam.max_marks) * 100, 1) if marks is not None else None
        subjects.append({
            "subject_name": subject_name,
            "term": exam.term,
            "exam_type": exam.exam_type,
            "marks": marks,
            "max_marks": exam.max_marks,
            "percentage": pct,
            "grade": score.grade if score else None,
        })
        term_key = str(exam.term)
        exam_type_scores.setdefault(term_key, {})
        if exam.exam_type not in exam_type_scores[term_key]:
            exam_type_scores[term_key][exam.exam_type] = []
        if marks is not None:
            exam_type_scores[term_key][exam.exam_type].append((marks / exam.max_marks) * 100)

    term_averages = {}
    all_pcts = []
    for term_key, exam_types in exam_type_scores.items():
        term_pcts = []
        for etype_pcts in exam_types.values():
            term_pcts.extend(etype_pcts)
        if term_pcts:
            avg = round(sum(term_pcts) / len(term_pcts), 1)
            term_averages[term_key] = avg
            all_pcts.extend(term_pcts)

    overall_average = round(sum(all_pcts) / len(all_pcts), 1) if all_pcts else 0.0

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    elements = []

    title_style = ParagraphStyle("Title2", parent=styles["Title"], fontSize=16, spaceAfter=6)
    subtitle_style = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=11, alignment=1, spaceAfter=4)
    body_style = styles["Normal"]
    center_style = ParagraphStyle("Center", parent=body_style, fontSize=11, alignment=1, spaceAfter=4)
    small_style = ParagraphStyle("Small", parent=body_style, fontSize=9, spaceAfter=2)

    school_name = school.name if school else "School"
    school_code = school.code if school else ""
    elements.append(Paragraph(f"{school_name}", title_style))
    elements.append(Paragraph(f"School Code: {school_code}", subtitle_style))
    elements.append(Spacer(1, 0.4 * cm))

    gender_display = (learner.gender or "-").upper()
    info_data = [
        [Paragraph("<b>Learner Name:</b>", body_style), Paragraph(learner.full_name, body_style),
         Paragraph("<b>Admission No:</b>", body_style), Paragraph(learner.admission_no or "-", body_style)],
        [Paragraph("<b>Gender:</b>", body_style), Paragraph(gender_display, body_style),
         Paragraph("<b>Class:</b>", body_style), Paragraph(school_class.name, body_style)],
        [Paragraph("<b>Academic Year:</b>", body_style), Paragraph(academic_year, body_style),
         Paragraph("", body_style), Paragraph("", body_style)],
    ]
    info_table = Table(info_data, colWidths=[3.5 * cm, 5.5 * cm, 3.5 * cm, 5.5 * cm])
    info_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 0.5 * cm))

    terms = sorted({s["term"] for s in subjects})
    for term in terms:
        term_subjects = [s for s in subjects if s["term"] == term]
        elements.append(Paragraph(f"<b>Term {term}</b>", center_style))
        elements.append(Spacer(1, 0.2 * cm))

        header = [
            Paragraph("<b>Subject</b>", body_style),
            Paragraph("<b>Exam Type</b>", body_style),
            Paragraph("<b>Marks</b>", body_style),
            Paragraph("<b>Max</b>", body_style),
            Paragraph("<b>%</b>", body_style),
            Paragraph("<b>Grade</b>", body_style),
        ]
        table_data = [header]
        for s in term_subjects:
            table_data.append([
                Paragraph(s["subject_name"], body_style),
                Paragraph(s["exam_type"], body_style),
                Paragraph(str(s["marks"]) if s["marks"] is not None else "-", body_style),
                Paragraph(str(s["max_marks"]), body_style),
                Paragraph(f"{s['percentage']}%" if s["percentage"] is not None else "-", body_style),
                Paragraph(s["grade"] or "-", body_style),
            ])

        col_widths = [6 * cm, 3.5 * cm, 2.2 * cm, 2 * cm, 2 * cm, 2 * cm]
        term_table = Table(table_data, colWidths=col_widths, repeatRows=1)
        term_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
        ]))
        elements.append(term_table)
        elements.append(Spacer(1, 0.4 * cm))

        term_avg = term_averages.get(str(term))
        if term_avg is not None:
            avg_style = ParagraphStyle("Avg", parent=body_style, fontSize=10, alignment=2, spaceAfter=4)
            elements.append(Paragraph(f"<b>Term {term} Average: {term_avg}%</b>", avg_style))
        elements.append(Spacer(1, 0.3 * cm))

    pct_style = ParagraphStyle("Pct", parent=body_style, fontSize=12, alignment=1, spaceAfter=6)
    elements.append(Paragraph(f"<b>Overall Average: {overall_average}%</b>", pct_style))
    elements.append(Spacer(1, 0.5 * cm))

    grade_descriptors = [
        "A (80-100%): Exceeding expectation",
        "B (65-79%): Meeting expectation",
        "C (50-64%): Approaching expectation",
        "D (30-49%): Below expectation",
        "E (0-29%): Far below expectation",
    ]
    elements.append(Paragraph("<b>Grade Descriptors</b>", ParagraphStyle("Section", parent=body_style, fontSize=12, spaceAfter=4)))
    for desc in grade_descriptors:
        elements.append(Paragraph(desc, small_style))
    elements.append(Spacer(1, 0.3 * cm))

    elements.append(Spacer(1, 0.5 * cm))
    sig_style = ParagraphStyle("Sig", parent=body_style, fontSize=10)
    elements.append(Paragraph("_" * 40 + "          " + "_" * 20, sig_style))
    elements.append(Paragraph("Class Teacher Signature                Date", sig_style))
    elements.append(Spacer(1, 0.4 * cm))
    elements.append(Paragraph("_" * 40 + "          " + "_" * 20, sig_style))
    elements.append(Paragraph("Principal Signature                    Date", sig_style))

    doc.build(elements)
    buffer.seek(0)
    data = buffer.read()
    filename = f"sba_report_card_{learner.full_name.replace(' ', '_')}_{academic_year}.pdf"
    result = _make_result(data, filename, "application/pdf")
    return data, result


async def _build_class_summary_csv(session, payload: dict) -> dict:
    import csv

    from app.models.assessment import Assessment
    from app.models.learner import Learner
    from app.models.run import AssessmentRun
    from app.models.school_class import SchoolClass
    from app.models.score import Score

    class_id = UUID(payload["class_id"])
    run_id = UUID(payload["run_id"])
    school_id = UUID(payload["school_id"])

    run = (
        await session.execute(
            select(AssessmentRun).where(AssessmentRun.id == run_id, AssessmentRun.school_id == school_id)
        )
    ).scalar_one_or_none()
    if run is None:
        raise ValueError("Run not found")

    assessment = (
        await session.execute(
            select(Assessment).where(
                Assessment.id == run.assessment_id,
                Assessment.school_id == school_id,
                Assessment.deleted_at.is_(None),
            )
        )
    ).scalar_one_or_none()
    if assessment is None:
        raise ValueError("Assessment not found")

    school_class = (
        await session.execute(
            select(SchoolClass).where(SchoolClass.id == class_id, SchoolClass.school_id == school_id)
        )
    ).scalar_one_or_none()
    if school_class is None:
        raise ValueError("Class not found")

    if school_class.id != class_id:
        raise ValueError("Class does not match run")

    learners = (
        await session.execute(
            select(Learner).where(
                Learner.class_id == class_id,
                Learner.school_id == school_id,
                Learner.deleted_at.is_(None),
            ).order_by(Learner.full_name)
        )
    ).scalars().all()
    if not learners:
        raise ValueError("No learners found in this class")

    learner_ids = [l.id for l in learners]
    scores = (
        await session.execute(
            select(Score).where(Score.run_id == run_id, Score.learner_id.in_(learner_ids)).order_by(Score.learner_id, Score.item_id)
        )
    ).scalars().all()

    scores_by_learner = {}
    for s in scores:
        scores_by_learner.setdefault(s.learner_id, {})[s.item_id] = s.level

    item_ids = [item.get("id", "") for item in (assessment.items or [])]
    seen = set(item_ids)
    for sid_map in scores_by_learner.values():
        for iid in sid_map:
            if iid not in seen:
                item_ids.append(iid)
                seen.add(iid)

    output = io.StringIO()
    writer = csv.writer(output)
    header = ["Learner Name", "Admission No"] + [f"Item {i + 1}" for i in range(len(item_ids))] + ["Total %"]
    writer.writerow(header)

    max_possible = sum(
        next((item.get("max_level", 4) for item in (assessment.items or []) if item.get("id") == iid), 4)
        for iid in item_ids
    )

    for learner in learners:
        row_scores = scores_by_learner.get(learner.id, {})
        total = 0
        item_values = []
        for iid in item_ids:
            level = row_scores.get(iid)
            item_values.append(str(level) if level is not None else "")
            if level is not None:
                total += level
        pct = round((total / max_possible) * 100, 1) if max_possible > 0 else 0.0
        writer.writerow([learner.full_name, learner.admission_no or ""] + item_values + [pct])

    data = output.getvalue().encode("utf-8")
    filename = f"class_summary_{school_class.name.replace(' ', '_')}_{run_id}.csv"
    result = _make_result(data, filename, "text/csv")
    return data, result


async def _build_term_exam_class_csv(session, payload: dict) -> dict:
    import csv

    from app.models.learner import Learner
    from app.models.learner_exam_score import LearnerExamScore
    from app.models.school_class import SchoolClass
    from app.models.term_exam import TermExam

    class_id = UUID(payload["class_id"])
    academic_year = payload["academic_year"]
    school_id = UUID(payload["school_id"])

    exams = (
        await session.execute(
            select(TermExam).where(
                TermExam.class_id == class_id,
                TermExam.academic_year == academic_year,
                TermExam.school_id == school_id,
            ).order_by(TermExam.term, TermExam.exam_type)
        )
    ).scalars().all()
    if not exams:
        raise ValueError("No term exams found for this class in this academic year")

    learners = (
        await session.execute(
            select(Learner).where(
                Learner.class_id == class_id,
                Learner.school_id == school_id,
                Learner.deleted_at.is_(None),
            ).order_by(Learner.full_name)
        )
    ).scalars().all()
    if not learners:
        raise ValueError("No learners found in this class")

    exam_ids = [e.id for e in exams]
    scores = (
        await session.execute(
            select(LearnerExamScore).where(
                LearnerExamScore.learner_id.in_([l.id for l in learners]),
                LearnerExamScore.term_exam_id.in_(exam_ids),
            )
        )
    ).scalars().all()

    scores_by_learner = {}
    for s in scores:
        scores_by_learner.setdefault(s.learner_id, {})[s.term_exam_id] = s

    school_class = await session.get(SchoolClass, class_id)
    output = io.StringIO()
    writer = csv.writer(output)

    headers = ["Learner Name", "Admission No"] + [f"{e.term}-{e.exam_type}" for e in exams] + ["Overall Average"]
    writer.writerow(headers)

    for learner in learners:
        row = [learner.full_name, learner.admission_no or ""]
        total_pct = []
        for exam in exams:
            score = scores_by_learner.get(learner.id, {}).get(exam.id)
            if score:
                pct = round((score.marks / exam.max_marks) * 100, 1)
                row.append(f"{score.marks}/{exam.max_marks} ({pct}%)")
                total_pct.append(pct)
            else:
                row.append("-")
        avg = round(sum(total_pct) / len(total_pct), 1) if total_pct else 0.0
        row.append(f"{avg}%")
        writer.writerow(row)

    data = output.getvalue().encode("utf-8")
    filename = f"term_exam_{school_class.name.replace(' ', '_')}_{academic_year}.csv"
    result = _make_result(data, filename, "text/csv")
    return data, result
